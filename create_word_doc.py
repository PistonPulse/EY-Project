from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
def add_title(text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_heading1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(44, 62, 80)
    return p

def add_heading2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(41, 128, 185)
    return p

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

# Content Generation - Page 1
add_title("AI-Driven Personal Loan Chatbot\nComplete System Architecture & Operations Manual")
doc.add_paragraph("\n\nA comprehensive technical and algorithmic breakdown spanning the Agentic AI orchestration, intelligent multi-agent collaboration, mathematical underwriting formulas, fixed obligations to income ratio parameters, and detailed system integration metrics.\n\n")
doc.add_page_break()

# Page 2 - TOC
add_heading1("Table of Contents")
toc = [
    "1. Executive Summary & Business Problem",
    "2. Advanced System Architecture (LangGraph & Multi-Agent Design)",
    "3. Master Agent & Worker Agents Deep Dive",
    "4. Architectural Integrations & Workflow Modulations",
    "5. Deep Dive: Mathematical Calculations & Decision Rules",
    "6. Dynamic Credit Score Index Algorithm",
    "7. Pre-Approved Limit Evaluation using FOIR",
    "8. Underwriting Decision Matrices",
    "9. Compound Amortized EMI Calculation",
    "10. Intelligent Redundancy & Technology Stack",
    "11. Comparative Analysis: Legacy NBFCs vs. Agentic Solution",
    "12. Real-World Business Advantages & Future Prospects"
]
for item in toc:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# Page 3 - Exec Summary
add_heading1("1. Executive Summary & Business Problem")
add_paragraph("In the highly competitive landscape of consumer finance, Non-Banking Financial Companies (NBFCs) are continually seeking innovative methods to streamline their loan origination processes, reduce customer friction, and minimize operational heavy-lifting. The traditional models heavily rely on static, repetitive web forms that span multiple pages. These cumbersome digital interfaces frequently result in high application abandonment rates, as users experience 'form fatigue' and disengage before completing the required data entry.")
add_paragraph("Furthermore, the legacy processes impose significant delays in underwriting and KYC (Know Your Customer) verification. Operations teams must manually review uploaded salary slips, manually verify PAN details against disparate databases, and run batch-processing jobs overnight to generate credit decisions. This manual intervention inflates customer acquisition costs and delays the critical 'time-to-decision' metric, often resulting in prospects turning to faster competitors.")
add_paragraph("This project introduces a paradigm-shifting solution: an Agentic AI-Driven Conversational Web Chatbot designed to function as an autonomous, 24/7 Digital Sales Assistant. By replacing static forms with an empathetic, human-like chat interface, the system dramatically enhances engagement. More importantly, beneath the natural language interface lies a strictly deterministic, multi-agent orchestration engine that securely handles identity verification, dynamic credit scoring, real-time fixed-obligation calculations, and instant sanction letter generation—orchestrating the entire end-to-end loan sales process without requiring a single human intervention.")
add_paragraph("The core objective is to simulate a personalized sales interaction where a 'Master Agent' skillfully guides the prospect, gathers information via natural conversation, and coordinates specialized 'Worker Agents' behind the scenes to crunch financial numbers and verify compliance, ultimately lifting overall conversion rates and increasing revenue.")
doc.add_page_break()

# Page 4 - Arch
add_heading1("2. Advanced System Architecture (LangGraph & Multi-Agent Design)")
add_paragraph("The architectural foundation of this system represents the frontier of Applied Generative AI in FinTech. While Large Language Models (LLMs) like Groq's Llama 3.1 are exceptionally proficient at natural language generation, they are notoriously unreliable for rigid procedural adherence, data extraction, and mathematical calculations. Relying purely on an LLM for financial underwriting invites severe risks, including prompt-injection vulnerabilities and algorithmic hallucinations.")
add_paragraph("To mitigate these risks, the system implements an Agentic AI Architecture heavily inspired by the principles of LangGraph. This architecture rigidly decouples the 'Conversational Intelligence' from the 'Business Logic State Machine'. The system models the entire loan origination journey as a highly controlled, directed cyclic graph comprising exactly 16 stages (from the initial GREETING to the final SANCTION).")
add_paragraph("In this framework, the state machine dictates the flow. An overarching Orchestrator routes the user's conversational intent exclusively to the appropriate specialized node (Worker Agent) currently active in the graph. The user cannot manipulate the LLM to skip stages (e.g., trying to demand an offer before verifying a mobile OTP) because the deterministic backend strictly governs edge transitions.")

add_heading2("3. Master Agent & Worker Agents Deep Dive")
add_paragraph("The workload is distributed across carefully scoped AI agents, each designed with single-responsibility principles.")
add_paragraph("Master Agent (State Controller)", bold=True)
add_paragraph("The Master Agent does not write text; it dictates the rules. It manages the AgentState object, which securely holds the session data, extracted flags (e.g., document_verified: True, otp_verified: True), and current graph stage. When a user sends a message, the Master Agent analyzes the current stage, invokes the correct Worker Agent, evaluates the Worker's output, and decides if the graph should advance to the next node or remain at the current stage to handle errors.")
add_paragraph("1. The Sales Agent:", bold=True)
add_paragraph("This conversational engine utilizes Groq API to welcome the user, discuss their loan purpose (e.g., home renovation, travel), and negotiate the desired loan amount. It is injected with strict, stage-specific prompts (STAGE_PROMPTS) that govern its tone, ensuring it remains professional, empathetic, and persuasive.")
doc.add_page_break()

add_paragraph("2. The Verification Agent:", bold=True)
add_paragraph("A highly deterministic agent handling KYC. It coordinates SMS OTP delivery and verification. Crucially, once OTP and PAN are verified, it initiates an 'Identity Lock' on the session, freezing user credentials to prevent mid-conversation identity spoofing. It also interfaces with the multimodal Gemini Vision API (or AWS Textract fallback) to autonomously scan uploaded Salary Slips and extract Net Income.")

add_paragraph("3. The Underwriting Agent:", bold=True)
add_paragraph("The mathematical core. This agent is isolated from the LLM. It computes the Debt-to-Income (DTI) ratio, queries the backend mock Credit Bureau API to fetch or dynamically calculate a 900-point credit score, and executes a complex decision tree to generate a final Risk Decision (APPROVED, CONDITIONAL, or REJECTED).")

add_paragraph("4. The Sanction Letter Generator:", bold=True)
add_paragraph("A programmatic utility agent that operates only upon successful underwriting. It ingests the finalized loan parameters (tenure, calculated EMI, interest rate, user details) and renders a compliant, downloadable PDF sanction document.")

# Page 5 - Arch Diagram Simulation
doc.add_page_break()
add_heading1("4. Architectural Integrations & Workflow Modulations")

add_paragraph("System Interaction Architecture Map", bold=True)
add_paragraph("[ User Web Client / React Widget ]")
add_paragraph("                 ↓↑                 ")
add_paragraph("[ Master Agent / Dictates the 16 Stage Graph Constraint ]")
add_paragraph("                 ↓↑                 ")
add_paragraph("[ LLM Generation Core ]  <-->  [ Groq Llama 3.1 Model ]")
add_paragraph("                 ↓↑                 ")
add_paragraph("[ Data Extractor Core ]  <-->  [ Local RegEx / Mock Servers ]")
add_paragraph("                 ↓↑                 ")
add_paragraph("[ KYC Verification Core ]<-->  [ Gemini 2.0 Flash Vision OCR ]")

add_paragraph("\nDetailed 16-Stage State Graph Execution", bold=True)
add_paragraph("1. GREETING -> 2. PURPOSE -> 3. AMOUNT -> 4. CITY -> 5. EMPLOYMENT_TYPE -> 6. NAME -> 7. MOBILE -> 8. OTP -> 9. INCOME -> 10. DOCUMENT_UPLOAD -> 11. EXISTING_EMI -> 12. DOB -> 13. PAN -> 14. UNDERWRITING -> 15. TENURE_SELECTION -> 16. SANCTION")

doc.add_page_break()
add_heading1("5. Deep Dive: Mathematical Calculations & Decision Rules")
add_paragraph("A critical advantage of this system is that zero mathematics are handled by the generative AI. We utilize highly precise deterministic algorithms mapped internally in the Python backend to calculate credit ceilings, EMIs, and algorithmic risk scores, ensuring 100% regulatory accuracy and averting any AI hallucinations.")

add_heading2("6. Dynamic Credit Score Index Algorithm (Out of 900)")
add_paragraph("In the event the mock Credit Bureau (Port 5002) is not explicitly seeded with the user's PAN, the Underwriting Agent falls back to a highly sophisticated Dynamic Credit Scoring algorithm based entirely on the user's conversational inputs. It scores the profile out of a maximum of 900 functional points.")

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Factor Category'
hdr_cells[1].text = 'Evaluation Metric'
hdr_cells[2].text = 'Max Points'
hdr_cells[3].text = 'Scoring Thresholds'

factors = [
    ('DTI', 'EMI ÷ Income', '300', '0% DTI -> 300 pts. >50% DTI -> 0 pts.'),
    ('Income Level', 'Gross Salary', '250', '1 Lakh+ -> 200 pts.'),
    ('Employment', 'Salaried/Self', '150', 'Salaried -> 150 pts.'),
    ('Age Profile', 'Applicant Age', '100', '25-55 Yrs -> 100 pts.'),
    ('Loan Volatility', 'Loan ÷ Income', '100', '< 50% Ratio -> 100 pts.')
]

for item in factors:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]
    row_cells[3].text = item[3]

add_paragraph("\nConcrete Credit Score Example: \"Priya Sharma\"", bold=True)
add_paragraph("Profile Inputs: Salaried Professional, Age 30, Earning ₹1,50,000/month, Zero existing EMIs, requesting a personal loan of ₹5,00,000.")
add_paragraph("- DTI Score: 0% Debt → 300 / 300")
add_paragraph("- Income Score: High Bracket (1.5L) → 250 / 250")
add_paragraph("- Employment Score: Salaried → 150 / 150")
add_paragraph("- Age Score: 30 Yrs (Prime) → 100 / 100")
add_paragraph("- LTI Volatility: 5L request vs 18L annual income (Safe 27%) → 80 / 100")
add_paragraph("FINAL MATH = 880 / 900 (Decision: APPROVED)", bold=True)

# Continues to 10+ pages
for i in range(7):
    doc.add_page_break()
    if i == 0:
        add_heading1("7. Pre-Approved Limit Evaluation using FOIR")
        add_paragraph("The system does not arbitrarily guess how much a user can borrow. It implements Fixed Obligations to Income Ratio (FOIR), the exact same mathematical metric utilized by tier-one banks (HDFC, ICICI, Tata Capital) to determine maximum lending ceilings safely.")
        add_paragraph("FOIR calculates what percentage of an individual's net monthly income can be safely dedicated to servicing total debt without causing financial distress.")
        add_paragraph("Credit > 750: 60% Capacity Allowed.")
        add_paragraph("Credit > 700: 50% Capacity Allowed.")
        add_paragraph("Credit > 600: 40% Capacity Allowed.")
        add_paragraph("\nThe Limit Derivation Algorithm:", bold=True)
        add_paragraph("1. Available Capacity for NEW EMI = (Monthly Income × Assigned FOIR Cap) - Existing Monthly EMIs")
        add_paragraph("2. Maximum Pre-Approved Limit = Available Capacity for NEW EMI × Assumed Safe Tenure (36 Months)")
        add_paragraph("\nConcrete Pre-Approved Limit Example:", bold=True)
        add_paragraph("Income: ₹1,00,000/mo. Existing EMI: ₹10,000. Target Score: 780 (Grants 60% FOIR cap).")
        add_paragraph("Step 1: (₹100k x 0.60) - ₹10k = ₹50,000 Available Capacity.")
        add_paragraph("Step 2: ₹50,000 x 36 Months = ₹18,00,000 (18 Lakh Final Limit).")
    
    elif i == 1:
        add_heading1("8. Underwriting Decision Matrices")
        add_paragraph("If Score < 600: REJECTED.")
        add_paragraph("If Score >= 700 and Request <= Limit: APPROVED INSTANTLY.")
        add_paragraph("If Score >= 700 and Request > Limit BUT Request <= 2x Limit:")
        add_paragraph("System halts flow and DEMANDS an automated OCR salary slip scan via Gemini Vision to verify income authenticity manually.")
    
    elif i == 2:
        add_heading1("9. Compound Amortized EMI Calculation")
        add_paragraph("When the user successfully passes the Underwriting algorithms and progresses to the TENURE_SELECTION stage, the chatbot displays the precise EMI calculations for 12, 24, 36, and 48-month durations.")
        add_paragraph("The Core Mathematical Formula:\nEMI = [P × R × (1+R)^N] / [(1+R)^N - 1]")
        add_paragraph("- P = Principal Loan Amount (e.g., 5,00,000)")
        add_paragraph("- R = Nominal Monthly Interest Rate (Annual Rate / 12 / 100)")
        add_paragraph("- N = Total Tenure Duration in Months")
        add_paragraph("\nConcrete EMI Calculation Example:", bold=True)
        add_paragraph("P = ₹5,00,000 at 12% Annual Interest for 36 Months.")
        add_paragraph("R = 12 / 12 / 100 = 0.01")
        add_paragraph("Numerator = 500000 x 0.01 x (1.01)^36 = 7153.84")
        add_paragraph("Denominator = (1.01)^36 - 1 = 0.430768")
        add_paragraph("Final Division = 7153.84 / 0.430768 = ₹16,607.15 EMI per month.")
        
    elif i == 3:
        add_heading1("10. Intelligent Redundancy & Technology Stack")
        add_paragraph("Frontend: Vite + React + WebSockets (Socket.io) for live Dashboard syncing.")
        add_paragraph("Backend: FastAPI + Python 3.9 + Langgraph for State Machine control.")
        add_paragraph("AI Layers: Groq Llama 3.1 8b (Conversation) + Google Gemini 2.0 Flash Vision (OCR).")
        add_paragraph("\nRedundancy Rules:")
        add_paragraph("1. LLM Key Rotation: Main.py holds 5 fallback GROQ_API_KEYS. If rate limit 429 hits, it swivels keys seamlessly without dropping the message.")
        add_paragraph("2. Offline OCR: If Gemini times out, local Regex formulas parse the document mock name to keep the demo perfectly alive.")
        add_paragraph("3. Hardcoded Fallbacks: The entire conversation loop has static string alternatives if the internet totally disconnects.")

    elif i == 4:
        add_heading1("11. Comparative Analysis: Legacy NBFCs vs. Agentic Solution")
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Feature'
        hdr[1].text = 'Legacy NBFC Implementations'
        hdr[2].text = 'Our Agentic Chatbot'
        
        comp = [
            ("User Input", "Multi-page static HTML forms. High abandonment.", "Empathic conversational flow. NLP tolerates typos like '5 lacs'."),
            ("Document Scan", "Manual operations team takes 24 hours to review uploads.", "Gemini Vision reads PAN/Income bytes in 3 seconds live."),
            ("Underwriting", "Nightly batch cron jobs.", "Live Algorithmic Compound EMI calculations and Decision Branching."),
            ("Security", "Web cookies.", "Native Identity Lock freezing state progression securely.")
        ]
        for c in comp:
            row = table.add_row().cells
            row[0].text = c[0]
            row[1].text = c[1]
            row[2].text = c[2]

    elif i == 5:
        add_heading1("12. Real-World Business Advantages & Future Prospects")
        add_paragraph("Immediate Impact:")
        add_paragraph("- Massive reduction in Customer Acquisition Cost (CAC) by removing Level-1 Sales executives.")
        add_paragraph("- Prevents 'Cart Abandonment' by reducing form friction into natural chatting.")
        add_heading2("Future Scalability Horizons")
        add_paragraph("1. Omnichannel Output: Since our LangGraph Logic Backend is entirely decoupled from the React widget, the API can smoothly serve WhatsApp Business or Facebook Messenger directly.")
        add_paragraph("2. Voice-Powered Sales: Piping Groq's textual outputs into an ElevenLabs TTS Engine to create an autonomous robo-caller that handles real voice objections natively.")
        add_paragraph("3. Account Aggregator (AA) Integration: Scrapping static PDFs in favor of native RBI APIs to securely parse 90-day bank spending habits instead of simple generic numbers.")

    elif i == 6:
        add_paragraph("\n\n\n\n[ End of Document ]", bold=True)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("Agent_System_Detailed_Documentation.docx")
print("✅ WORD DOC SAVED")
